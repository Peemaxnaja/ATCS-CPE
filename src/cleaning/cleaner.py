import logging
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Union
import yaml
from pydantic import BaseModel, Field

from src.collection.collector import CollectionOutput
from src.utils.base_step import PipelineStep

# Try importing parsing libraries
try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None


# Logging Configuration
logger = logging.getLogger(__name__)
if not logger.handlers:
    handler = logging.StreamHandler()
    formatter = logging.Formatter("[%(asctime)s] [%(levelname)s] [%(name)s]: %(message)s")
    handler.setFormatter(formatter)
    logger.addHandler(handler)
    logger.setLevel(logging.INFO)


class CleaningOutput(BaseModel):
    """Data Contract output for Cleaning Step."""
    clean_files_directory: str = Field(..., description="Absolute path of the clean text files directory")
    cleaned_file_list: List[str] = Field(default_factory=list, description="List of absolute paths of cleaned txt files")


class Cleaner(PipelineStep):
    """
    Step 2: Cleaning Module
    Extracts raw text from multiple file formats (.pdf, .docx, .html, .txt, .md),
    removes noise (HTML tags, headers/footers, excess whitespace, navigation),
    and saves plain text files (.txt) into data/clean/.
    """

    def __init__(self, config_path: Union[str, Path] = "config/config.yaml") -> None:
        self.config_path = Path(config_path)
        self.config = self._load_config()

    def _load_config(self) -> Dict[str, Any]:
        """Loads configuration from YAML file or returns defaults."""
        if self.config_path.exists():
            try:
                with open(self.config_path, "r", encoding="utf-8") as f:
                    config = yaml.safe_load(f)
                    if config and "cleaning" in config:
                        return config["cleaning"]
            except Exception as e:
                logger.warning(f"Failed to load config from {self.config_path}: {e}. Using defaults.")

        return {
            "raw_directory": "data/raw",
            "clean_directory": "data/clean",
            "remove_extra_whitespace": True,
            "min_text_length": 5,
        }

    # --- File Extractors ---

    def _extract_pdf(self, file_path: Path) -> str:
        """Extracts plain text from PDF files using pypdf."""
        if pypdf is None:
            logger.warning("pypdf library is missing; skipping PDF text extraction.")
            return ""
        
        extracted_text = []
        try:
            reader = pypdf.PdfReader(file_path)
            for page_idx, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                extracted_text.append(text)
            return "\n".join(extracted_text)
        except Exception as e:
            logger.error(f"Error reading PDF file '{file_path}': {e}")
            return ""

    def _extract_docx(self, file_path: Path) -> str:
        """Extracts plain text from Word (.docx) files using python-docx."""
        if docx is None:
            logger.warning("python-docx library is missing; skipping DOCX text extraction.")
            return ""
        
        try:
            doc = docx.Document(file_path)
            full_text = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        full_text.append(row_text)
            return "\n".join(full_text)
        except Exception as e:
            logger.error(f"Error reading DOCX file '{file_path}': {e}")
            return ""

    def _extract_html(self, file_path: Path) -> str:
        """Extracts text content from HTML files using BeautifulSoup, removing noise elements."""
        try:
            content = file_path.read_text(encoding="utf-8", errors="ignore")
            if BeautifulSoup is None:
                # Basic fallback HTML tag stripping regex
                return re.sub(r"<[^>]+>", " ", content)
            
            soup = BeautifulSoup(content, "html.parser")

            # Decompose boilerplate & noise elements
            for tag in soup(["script", "style", "nav", "header", "footer", "aside", "form"]):
                tag.decompose()

            # Extract clean text with newline separators
            return soup.get_text(separator="\n")
        except Exception as e:
            logger.error(f"Error parsing HTML file '{file_path}': {e}")
            return ""

    def _extract_markdown(self, file_path: Path) -> str:
        """Extracts plain text from Markdown (.md) files, removing formatting artifacts."""
        try:
            text = file_path.read_text(encoding="utf-8", errors="ignore")
            # Remove images ![alt](url)
            text = re.sub(r"!\[.*?\]\(.*?\)", "", text)
            # Replace links [label](url) with label
            text = re.sub(r"\[([^\]]+)\]\(.*?\)", r"\1", text)
            # Remove code block indicators ```
            text = re.sub(r"```[a-zA-Z]*", "", text)
            # Remove headers markers (#)
            text = re.sub(r"^#+\s*", "", text, flags=re.MULTILINE)
            # Remove horizontal rules
            text = re.sub(r"^\s*[-*_]{3,}\s*$", "", text, flags=re.MULTILINE)
            # Remove bold/italic emphasis
            text = re.sub(r"[*_]{1,3}([^*_]+)[*_]{1,3}", r"\1", text)
            return text
        except Exception as e:
            logger.error(f"Error reading Markdown file '{file_path}': {e}")
            return ""

    def _extract_txt(self, file_path: Path) -> str:
        """Extracts text from plain TXT files."""
        try:
            return file_path.read_text(encoding="utf-8", errors="ignore")
        except Exception as e:
            logger.error(f"Error reading TXT file '{file_path}': {e}")
            return ""

    # --- Noise Cleaning Methods ---

    def clean_text(self, text: str) -> str:
        """
        Cleans raw extracted text by removing unwanted noise, page markers, 
        and excessive whitespaces.
        """
        if not text:
            return ""

        # 1. Remove common header/footer page numbers (e.g. "Page 1 of 5", "Page 1")
        text = re.sub(r"(?i)\bpage\s+\d+(\s+of\s+\d+)?\b", "", text)

        # 2. Trim trailing spaces per line
        lines = [line.strip() for line in text.splitlines()]

        # 3. Filter out repetitive blank lines or trivial divider lines
        cleaned_lines = []
        for line in lines:
            # Skip empty lines if previous line was also empty
            if line:
                cleaned_lines.append(line)
            elif cleaned_lines and cleaned_lines[-1] != "":
                cleaned_lines.append("")

        result = "\n".join(cleaned_lines)

        # 4. Collapse multiple blank lines (3+ newlines to 2)
        result = re.sub(r"\n{3,}", "\n\n", result)

        return result.strip()

    # --- Execution Flow ---

    def execute(self, input_data: Optional[Union[CollectionOutput, Dict[str, Any], str, Path]] = None) -> CleaningOutput:
        """
        Executes the file cleaning process.
        
        :param input_data: CollectionOutput object, dict, or directory path.
        :return: CleaningOutput containing clean directory path and list of cleaned text files.
        """
        logger.info("Starting Cleaning Step...")

        clean_dir_str = self.config.get("clean_directory", "data/clean")
        raw_dir_str = self.config.get("raw_directory", "data/raw")
        clean_path = Path(clean_dir_str).resolve()
        clean_path.mkdir(parents=True, exist_ok=True)

        files_to_clean: List[Path] = []

        # Determine target raw files to process
        if isinstance(input_data, CollectionOutput):
            files_to_clean = [Path(p) for p in input_data.file_list]
        elif isinstance(input_data, dict) and "file_list" in input_data:
            files_to_clean = [Path(p) for p in input_data["file_list"]]
        elif isinstance(input_data, (str, Path)):
            target_path = Path(input_data)
            if target_path.is_file():
                files_to_clean = [target_path]
            elif target_path.is_dir():
                files_to_clean = [p for p in target_path.glob("*") if p.is_file()]
        else:
            # Fallback to scanning data/raw
            raw_path = Path(raw_dir_str).resolve()
            if raw_path.exists():
                files_to_clean = [p for p in raw_path.glob("*") if p.is_file()]

        logger.info(f"Target clean directory: {clean_path}")
        logger.info(f"Found {len(files_to_clean)} raw files to clean.")

        cleaned_file_paths: List[str] = []

        for raw_file in files_to_clean:
            if not raw_file.exists() or raw_file.stat().st_size == 0:
                logger.warning(f"Skipping invalid/empty file: {raw_file}")
                continue

            suffix = raw_file.suffix.lower()
            extracted_text = ""

            if suffix == ".pdf":
                extracted_text = self._extract_pdf(raw_file)
            elif suffix == ".docx":
                extracted_text = self._extract_docx(raw_file)
            elif suffix == ".html":
                extracted_text = self._extract_html(raw_file)
            elif suffix == ".md":
                extracted_text = self._extract_markdown(raw_file)
            elif suffix == ".txt":
                extracted_text = self._extract_txt(raw_file)
            else:
                logger.warning(f"Unsupported file extension '{suffix}' for file '{raw_file.name}'")
                continue

            # Clean extracted text
            cleaned_text = self.clean_text(extracted_text)

            min_len = self.config.get("min_text_length", 5)
            if len(cleaned_text) < min_len:
                logger.warning(f"Cleaned text for '{raw_file.name}' is shorter than min_length ({min_len}). Skipping save.")
                continue

            # Save clean plain text file
            clean_file_name = f"{raw_file.stem}.txt"
            output_file_path = clean_path / clean_file_name
            
            try:
                output_file_path.write_text(cleaned_text, encoding="utf-8")
                cleaned_file_paths.append(str(output_file_path.resolve()))
                
                orig_size = len(extracted_text)
                clean_size = len(cleaned_text)
                reduction = ((orig_size - clean_size) / orig_size * 100) if orig_size > 0 else 0
                
                logger.info(
                    f"Successfully cleaned: '{raw_file.name}' -> '{clean_file_name}' "
                    f"({clean_size} chars, reduction: {reduction:.1f}%)"
                )
            except Exception as e:
                logger.error(f"Failed to write cleaned file '{output_file_path}': {e}")

        logger.info(f"Cleaning finished. Total cleaned files saved: {len(cleaned_file_paths)}")

        return CleaningOutput(
            clean_files_directory=str(clean_path),
            cleaned_file_list=cleaned_file_paths
        )
